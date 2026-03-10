import telebot
from telebot import types
import os
from docx import Document
from docx.shared import Pt

# Ваш токен
TOKEN = '8556430635:AAE2TYMC4mnjFOzWzitd8i3sxvuBPwDSfp8'
bot = telebot.TeleBot(TOKEN)

# Папка для временного хранения загруженных работ
if not os.path.exists('uploads'):
    os.makedirs('uploads')

# --- КЛАВИАТУРЫ ---
def get_main_menu():
    markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
    btn1 = types.KeyboardButton('📚 Требования')
    btn2 = types.KeyboardButton('🔍 Проверить работу')
    markup.add(btn1, btn2)
    return markup

# --- ОБРАБОТЧИКИ КОМАНД ---
@bot.message_handler(commands=['start'])
def send_welcome(message):
    text = (
        "Привет! Я бот для нормоконтроля.\n\n"
        "Я помогу проверить вашу курсовую или ВКР на соответствие основным требованиям оформления.\n"
        "Выберите нужное действие в меню ниже 👇"
    )
    bot.send_message(message.chat.id, text, reply_markup=get_main_menu())

@bot.message_handler(content_types=['text'])
def handle_text(message):
    if message.text == '📚 Требования':
        text = (
            "📌 **Основные требования:**\n"
            "• Формат: А4, поля по 20 мм.\n"
            "• Шрифт: Times New Roman, 14 пт, черный.\n"
            "• Интервал: 1.2, абзацный отступ: 1.0 см.\n\n"
            "🔗 [Ссылка на полные методические рекомендации](https://drive.google.com/drive/folders/19Mw0rgJdEjVy8DrLx8v74MiwZHIPhAh1?hl=ru)"
        )
        # Здесь также можно отправлять файл чек-листа
        bot.send_message(message.chat.id, text, parse_mode='Markdown')
        
    elif message.text == '🔍 Проверить работу':
        bot.send_message(message.chat.id, "📎 Пожалуйста, отправьте файл вашей работы в формате **.docx** (не более 50 Мб).", parse_mode='Markdown')
    else:
        bot.send_message(message.chat.id, "Используйте кнопки меню для взаимодействия.", reply_markup=get_main_menu())

# --- ОБРАБОТКА ДОКУМЕНТОВ ---
@bot.message_handler(content_types=['document'])
def handle_document(message):
    file_name = message.document.file_name
    file_size = message.document.file_size

    # Валидация
    if not file_name.endswith('.docx'):
        bot.reply_to(message, "❌ Ошибка: Поддерживаются только файлы формата .docx. Формат .doc или другие не подходят.")
        return
    
    if file_size > 50 * 1024 * 1024: # 50 MB
        bot.reply_to(message, "❌ Ошибка: Размер файла превышает 50 Мб.")
        return

    # Индикация процесса
    msg = bot.reply_to(message, "⏳ Идет загрузка и проверка документа, это может занять пару минут...")

    try:
        # Скачивание файла
        file_info = bot.get_file(message.document.file_id)
        downloaded_file = bot.download_file(file_info.file_path)
        
        file_path = os.path.join('uploads', file_name)
        with open(file_path, 'wb') as new_file:
            new_file.write(downloaded_file)
            
        # Запуск анализа (Ядро системы)
        report = analyze_docx(file_path)
        
        # Отправка отчета
        bot.edit_message_text(chat_id=message.chat.id, message_id=msg.message_id, text=report, parse_mode='Markdown')
        
        # Удаление временного файла
        os.remove(file_path)
        
    except Exception as e:
        bot.edit_message_text(chat_id=message.chat.id, message_id=msg.message_id, text=f"⚠️ Произошла ошибка при обработке файла: {e}")

# --- ЯДРО ПРОВЕРКИ (АНАЛИЗАТОР) ---
def analyze_docx(file_path):
    """
    Функция анализирует .docx файл на соответствие правилам из ТЗ.
    Возвращает строку с готовым отчетом.
    """
    try:
        doc = Document(file_path)
    except Exception:
        return "❌ Не удалось прочитать файл. Возможно, он поврежден."


    errors = []
    recommendations = []
    
    # 1. Проверка шрифта в абзацах (базовая реализация)
    font_errors = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.name and run.font.name != 'Times New Roman':
                font_errors += 1
                break # Считаем только первый проблемный кусок в абзаце
    
    if font_errors > 0:
        errors.append(f"• Найден шрифт, отличный от Times New Roman (в {font_errors} абзацах).")

    # 2. Проверка полей (базовая проверка первого раздела)
    try:
        section = doc.sections[0]
        # 360000 EMU = 1 см. 20 мм = 2 см = 720000 EMU
        left_margin_cm = round(section.left_margin.emu / 360000, 1)
        right_margin_cm = round(section.right_margin.emu / 360000, 1)
        
        if left_margin_cm != 2.0 or right_margin_cm != 2.0:
            errors.append(f"• Поля не соответствуют стандарту 20 мм (Текущие: Левое {left_margin_cm} см, Правое {right_margin_cm} см).")
    except:
        recommendations.append("• Не удалось определить размеры полей автоматически, проверьте их вручную.")

    # 3. Проверка слова "СОДЕРЖАНИЕ"
    content_found = False
    for para in doc.paragraphs:
        if "СОДЕРЖАНИЕ" in para.text.upper():
            content_found = True
            break
            
    if not content_found:
        errors.append("• Не найден раздел «СОДЕРЖАНИЕ».")

    # Формирование итогового отчета
    if not errors and not recommendations:
        return "✅ **Общий вывод: Ошибок не найдено!**\n\nВаша работа оформлена корректно по проверенным параметрам."
    
    report_text = "⚠️ **Общий вывод: Требуется доработка.**\n\n"
    
    if errors:
        report_text += "**Список найденных ошибок:**\n" + "\n".join(errors) + "\n\n"
        
    if recommendations:
        report_text += "**Рекомендации:**\n" + "\n".join(recommendations)
        
    return report_text

# Запуск бота
if __name__ == '__main__':
    print("Бот запущен...")
    bot.polling(none_stop=True)
